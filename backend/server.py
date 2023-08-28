from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt
import requests
import datetime
from flask_cors import CORS  # Import CORS

app = Flask(_name_)
CORS(app)

API_KEY = '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'

@app.route('/generate', methods=['GET'])
def generate():
    print("request sent")
    return jsonify({"error": "Internal Server Error"})

@app.route('/completions', methods=['POST'])
def completions():
    try:
        data = request.json
        options = {
            "model": "gpt-3.5-turbo",
            "messages": [{"role": "user", "content": data["messages"]}],
            "max_tokens": 1000,
        }
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json",
        }
        response = requests.post(
            "https://api.openai.com/v1/chat/completions", json=options, headers=headers
        )
        response_data = response.json()
        return jsonify(response_data)
    except Exception as e:
        print(e)
        return jsonify({"error": "Internal Server Error"}), 500

@app.route('/generate_document', methods=['POST'])
def generate_document():
    print("request sent")
    sy = request.json.get("input_text", "").lower()

    url = "http://localhost:8000/completions"
    data = {
        "messages": sy + " extract units and topics as array in this format [{unit:unit_name,topics:[topics..]}] strictly follow this format dont any other text on response use give the formatted output"
    }

    response = requests.post(url, json=data)
    result = response.json()
    arraystring = result['choices'][0]['message']['content']
    
    for x in range(len(arraystring)):
        if arraystring[x] == '[':
            array = arraystring[x:]
            break
    for x in range(len(array) - 1, -1, -1):
        if array[x] == ']':
            array = array[:x + 1]
            break

    content = eval(array)
    print(content)
    forpdf = []
    
    for j in range(len(content)):
        forunit = []
        for topic in content[j]["topics"]:
            res = requests.post(url, json={"messages": topic + " explain in 500 words with sub headings, dont mention count of words"}).json()
            contentt=res['choices'][0]['message']['content']
            # Assuming you have 'contentt' and 'topic' defined already
            content_lines = contentt.split('\n')
            if len(content_lines) > 0 and content_lines[0].lower() == topic.lower():
                content_lines.pop(0)  # Remove the first line if it matches 'topic' (case-insensitive)
            # Join the remaining lines back into 'contentt'
            contentt = '\n'.join(content_lines)
            forunit.append({"topic": topic, "content":  contentt})
            print(forunit)
        forpdf.append({"unit": content[j]["unit"], "topics": forunit})
    print(forpdf)
    print("Successs")
    
    # Create the Word document
    doc = Document()

    # Set font properties
    font_family = "Arial"
    font_size_unit_title = Pt(18)
    font_size_topic_title = Pt(12)
    font_size_content = Pt(12)

    # Function to add formatted content to the document
    def add_formatted_content(paragraph, text, font_size, bold=False):
        run = paragraph.add_run(text)
        font = run.font
        font.name = font_family
        font.size = font_size
        font.bold = bold

    # Add the content to the Word document
    for unit_data in forpdf:
        unit_title = unit_data["unit"]
        add_formatted_content(doc.add_paragraph(), unit_title, font_size_unit_title, bold=True)

        for topic_data in unit_data["topics"]:
            topic_title = topic_data["topic"]
            add_formatted_content(doc.add_paragraph(), topic_title, font_size_topic_title, bold=True)

            content = topic_data["content"]
            add_formatted_content(doc.add_paragraph(), content, font_size_content)
    print("sss")
    # Get the current date and time
    current_datetime = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    docx_file_path = f"Notes_{current_datetime}.docx"
    doc.save(docx_file_path)

    return send_file(docx_file_path, as_attachment=True)

if _name_ == '_main_':
    app.run(host='0.0.0.0', port=8000)